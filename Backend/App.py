from flask import Flask, request, jsonify, send_file
import time
import logging
import os
import json
import re
import base64
import random
import requests
from concurrent.futures import ThreadPoolExecutor, as_completed
from contextlib import contextmanager
from typing import Optional
from werkzeug.utils import secure_filename
import pymupdf as fitz
from google import genai
from google.genai import types as genai_types
from flask_cors import CORS
from tenacity import retry, stop_after_attempt, wait_fixed, retry_if_exception_type
from requests.exceptions import ConnectionError as RequestsConnectionError
from youtube_transcript_api import YouTubeTranscriptApi, NoTranscriptFound, TranscriptsDisabled
import datetime
import tempfile

# Use WARNING in production to reduce log noise and I/O overhead
_log_level = os.getenv("LOG_LEVEL", "WARNING").upper()
logging.basicConfig(level=getattr(logging, _log_level, logging.WARNING))
logger = logging.getLogger(__name__)

from dotenv import load_dotenv
env_path = os.path.join(os.path.dirname(__file__), '.env')
load_dotenv(env_path)
ENABLE_RAG = os.getenv("ENABLE_RAG", "1").lower() in ("1", "true", "yes")

RAGProcessor = None
rag_import_error = None
if ENABLE_RAG:
    try:
        from rag_processor import RAGProcessor
        logger.info("✓ RAG processor imported successfully")
    except Exception as e:
        rag_import_error = str(e)
        logger.warning(f"RAG processor import failed: {e}")
        RAGProcessor = None
else:
    rag_import_error = "RAG disabled (set ENABLE_RAG=1 to enable)"
    logger.info("RAG processor skipped (ENABLE_RAG=0)")

app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "*"}})

_max_mb = int(os.getenv("MAX_UPLOAD_MB", "20"))
app.config["MAX_CONTENT_LENGTH"] = _max_mb * 1024 * 1024

UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
GEMINI_CHAT_API_KEY = os.getenv("GEMINI_CHAT_API_KEY") or GEMINI_API_KEY
YOUTUBE_API_KEY = os.getenv("YOUTUBE_API_KEY")
GEMINI_MODEL = os.getenv("GEMINI_MODEL", "gemini-3.6-flash")

gemini_client = None
gemini_chat_client = None

if GEMINI_API_KEY:
    try:
        gemini_client = genai.Client(api_key=GEMINI_API_KEY)
        logger.info("✓ Gemini client initialized")
    except Exception as e:
        logger.warning(f"Gemini client init failed: {e}")
else:
    logger.warning("GEMINI_API_KEY not set")

# Reuse the same client if keys are identical (saves ~15 MB RAM)
if GEMINI_CHAT_API_KEY and GEMINI_CHAT_API_KEY != GEMINI_API_KEY:
    try:
        gemini_chat_client = genai.Client(api_key=GEMINI_CHAT_API_KEY)
        logger.info("✓ Separate Gemini chat client initialized")
    except Exception as e:
        logger.warning(f"Gemini chat client init failed: {e}")
else:
    gemini_chat_client = gemini_client

youtube = None
def _get_youtube_client():
    """Lazily build the YouTube Data API client on first use."""
    global youtube
    if youtube is not None:
        return youtube
    if not YOUTUBE_API_KEY:
        return None
    try:
        import googleapiclient.discovery
        youtube = googleapiclient.discovery.build(
            "youtube", "v3", developerKey=YOUTUBE_API_KEY
        )
        logger.info("✓ YouTube Data API client initialized")
    except Exception as e:
        logger.warning(f"YouTube Data API client failed: {e}")
        youtube = None
    return youtube

rag_processor = None
rag_init_error = None
if ENABLE_RAG and RAGProcessor is not None:
    try:
        backend_dir = os.path.dirname(os.path.abspath(__file__))
        persist_dir = os.path.join(backend_dir, 'vector_store')
        os.makedirs(persist_dir, exist_ok=True)
        rag_processor = RAGProcessor(persist_directory=persist_dir)
        logger.info(f"✓ RAG processor initialized with Gemini embeddings at: {persist_dir}")
    except Exception as e:
        rag_init_error = str(e)
        logger.error(f"RAG processor initialization failed: {e}", exc_info=True)
        rag_processor = None
else:
    if not ENABLE_RAG:
        logger.info("RAG processor disabled (set ENABLE_RAG=1 to enable)")
    elif rag_import_error:
        logger.warning(f"RAG processor not available: {rag_import_error}")


@app.errorhandler(413)
def request_too_large(e):
    max_mb = int(os.getenv("MAX_UPLOAD_MB", "20"))
    return jsonify({"error": f"File too large. Maximum upload size is {max_mb} MB. Please use a smaller file."}), 413


@app.get('/')
def index():
    return jsonify({"status": "ok", "service": "NoteLooms Backend"})


@app.get('/health')
def health():
    return jsonify({
        "status": "ok",
        "ai_initialized": gemini_client is not None,
        "rag_enabled": ENABLE_RAG,
        "rag_initialized": rag_processor is not None,
        "rag_backend": "gemini-embeddings" if rag_processor else None,
    })

def extract_video_id(url):
    """Extract video ID from various YouTube URL formats."""
    patterns = [
        r'(?:v=|\/)([0-9A-Za-z_-]{11}).*',
        r'(?:embed\/)([0-9A-Za-z_-]{11})',
        r'^([0-9A-Za-z_-]{11})$'
    ]
    for pattern in patterns:
        match = re.search(pattern, url)
        if match:
            return match.group(1)
    return None


def get_youtube_transcript_advanced(video_url_or_id):
    """Fetch transcript with multiple fallback options."""
    try:
        video_id = extract_video_id(video_url_or_id)
        if not video_id:
            logger.error(f"Could not extract video ID from: {video_url_or_id}")
            return None

        logger.info(f"Fetching transcript for video ID: {video_id}")

        try:
            transcript_data = YouTubeTranscriptApi.get_transcript(video_id, languages=['en'])
            formatted = [
                {'text': e['text'].strip(), 'start': float(e['start']), 'duration': float(e['duration'])}
                for e in transcript_data
            ]
            logger.info(f"✓ Fetched transcript (direct, {len(formatted)} segments)")
            return formatted
        except Exception as e:
            logger.warning(f"Direct transcript fetch failed: {e}")

        try:
            transcript_list = YouTubeTranscriptApi.list_transcripts(video_id)
            transcript = None
            try:
                transcript = transcript_list.find_manually_created_transcript(['en'])
                logger.info("✓ Found manually created English transcript")
            except Exception:
                pass
            if not transcript:
                try:
                    transcript = transcript_list.find_generated_transcript(['en'])
                    logger.info("✓ Found auto-generated English transcript")
                except Exception:
                    pass
            if transcript:
                data = transcript.fetch()
                formatted = [
                    {'text': e['text'].strip(), 'start': float(e['start']), 'duration': float(e['duration'])}
                    for e in data
                ]
                logger.info(f"✓ Fetched transcript (fallback, {len(formatted)} segments)")
                return formatted
        except TranscriptsDisabled:
            logger.warning(f"Transcripts disabled for video: {video_id}")
            return None
        except NoTranscriptFound:
            logger.warning(f"No transcripts found for video: {video_id}")
            return None
        except Exception as e:
            logger.error(f"Error getting transcript list: {e}")

        logger.error(f"✗ All transcript methods failed for video: {video_id}")
        return None

    except Exception as e:
        logger.error(f"Error fetching transcript: {e}")
        return None


def extract_text_from_youtube(url):
    """Extract text and transcript from YouTube video."""
    try:
        video_id = extract_video_id(url)
        if not video_id:
            return None, None

        transcript_data = get_youtube_transcript_advanced(url)
        extracted_text = None

        if transcript_data:
            extracted_text = " ".join(e["text"] for e in transcript_data)
            logger.info(f"✓ Extracted {len(transcript_data)} transcript segments")
        else:
            logger.warning(f"No transcript data for video: {video_id}")

        if not extracted_text:
            yt_client = _get_youtube_client()
            if yt_client:
                try:
                    resp = yt_client.videos().list(part="snippet", id=video_id).execute()
                    if resp.get("items"):
                        s = resp["items"][0]["snippet"]
                        extracted_text = f"Title: {s['title']}\nDescription: {s['description']}"
                        logger.info(f"✓ Got YouTube metadata for video: {video_id}")
                except Exception as e:
                    logger.error(f"YouTube metadata error: {e}")

        if not extracted_text:
            return None, None

        return extracted_text, transcript_data

    except Exception as e:
        logger.error(f"YouTube processing error: {e}", exc_info=True)
        return None, None


@app.route('/api/youtube/transcript', methods=['POST'])
def get_transcript_endpoint():
    data = request.json
    video_url = data.get('url')
    if not video_url:
        return jsonify({'success': False, 'error': 'No URL provided'}), 400

    transcript = get_youtube_transcript_advanced(video_url)
    video_id = extract_video_id(video_url)

    if transcript:
        return jsonify({'success': True, 'transcript': transcript, 'video_id': video_id})
    return jsonify({'success': False, 'error': 'Could not fetch transcript.', 'transcript': []})

@contextmanager
def open_pdf(pdf_path: str):
    """Context manager for safely opening PDF files."""
    doc = None
    try:
        doc = fitz.open(pdf_path)
        yield doc
    finally:
        if doc:
            doc.close()


def extract_text_from_pdf(pdf_path: str, max_pages: int = 100) -> Optional[str]:
    """Extract text from PDF using PyMuPDF with memory optimisation."""
    try:
        text_chunks = []
        with open_pdf(pdf_path) as doc:
            pages_to_process = min(len(doc), max_pages)
            for page_num in range(pages_to_process):
                try:
                    page_text = doc[page_num].get_text("text")
                    if page_text and page_text.strip():
                        text_chunks.append(f"--- Page {page_num + 1} ---\n{page_text.strip()}")
                    if page_num % 50 == 0:
                        gc.collect()
                except Exception as page_error:
                    logger.warning(f"Error processing page {page_num}: {page_error}")
        return "\n\n".join(text_chunks) if text_chunks else None
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
        return None


def extract_text_from_scanned_pdf(pdf_path: str, max_pages: int = 20, dpi: int = 200) -> Optional[str]:
    """Extract text from scanned PDF via OCR. Lazy-loads heavy deps only when needed."""
    try:
        from pdf2image import convert_from_path
        import cv2
        import numpy as np
        import pytesseract
    except ImportError as e:
        logger.error(f"OCR dependencies not available: {e}")
        return None

    with tempfile.TemporaryDirectory() as tmpdir:
        try:
            logger.info(f"Starting OCR for {pdf_path} (max {max_pages} pages)")
            images = convert_from_path(
                pdf_path, first_page=1, last_page=max_pages, dpi=dpi, output_folder=tmpdir
            )
            extracted_text = []
            for i, image in enumerate(images):
                if i >= max_pages:
                    break
                try:
                    img     = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)
                    gray    = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
                    denoised = cv2.medianBlur(gray, 3)
                    _, binary = cv2.threshold(denoised, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
                    custom_config = (
                        r"--oem 3 --psm 6 -c tessedit_char_whitelist="
                        r"ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789.,!?;:()[]{}@#$%^&*+-= "
                    )
                    text = pytesseract.image_to_string(binary, config=custom_config)
                    if text.strip():
                        extracted_text.append(f"--- Page {i + 1} ---\n{text.strip()}")
                    del img, gray, denoised, binary
                    if i % 5 == 0:
                        gc.collect()
                except Exception as img_e:
                    logger.error(f"Error processing page {i}: {img_e}")
            if extracted_text:
                logger.info(f"OCR extracted {len(extracted_text)} pages")
                return "\n\n".join(extracted_text)
            return None
        except Exception as e:
            logger.error(f"OCR error: {e}")
            return None

def _get_removal_delay(filepath: str) -> float:
    try:
        size_mb = os.path.getsize(filepath) / (1024 * 1024)
        if size_mb > 100: return 1.0
        if size_mb > 10:  return 0.5
        return 0.2
    except (OSError, TypeError):
        return 0.3


def safe_remove(filepath: str, max_retries: int = 3, initial_delay: float | None = None) -> bool:
    """Safely remove a file with retry + exponential back-off."""
    if not os.path.exists(filepath):
        return True
    delay = initial_delay if initial_delay is not None else _get_removal_delay(filepath)
    for attempt in range(max_retries):
        try:
            os.remove(filepath)
            return True
        except (PermissionError, OSError) as e:
            if attempt == max_retries - 1:
                logger.warning(f"Failed to remove {filepath}: {e}")
                return False
            time.sleep(min(delay * (2 ** attempt) + (random.random() * 0.1), 5.0))
    return False

def generate_gemini_response(prompt):
    """Generate response with proper error handling."""
    try:
        if not gemini_client:
            return "⚠ ERROR: Gemini client not initialized. Please check GEMINI_API_KEY."
        response = gemini_client.models.generate_content(model=GEMINI_MODEL, contents=prompt)
        text = getattr(response, "text", None)
        if not text or not text.strip():
            return "⚠ ERROR: No response generated."
        return text.strip().replace("*", "")
    except Exception as e:
        logger.error(f"Gemini API error: {e}")
        if "429" in str(e): return "⚠ ERROR: API quota exceeded. Please try again later."
        if "404" in str(e): return "⚠ ERROR: Model not available. Please check the model configuration."
        return f"⚠ ERROR: {e}"


def generate_chat_response(prompt):
    """Use dedicated chat API key if provided, falls back to main key."""
    try:
        if not gemini_chat_client:
            return "⚠ ERROR: Gemini chat client not initialized. Please check GEMINI_CHAT_API_KEY."
        response = gemini_chat_client.models.generate_content(model=GEMINI_MODEL, contents=prompt)
        text = getattr(response, "text", None)
        if not text or not text.strip():
            return "⚠ ERROR: No chat response generated."
        return text.strip().replace("*", "")
    except Exception as e:
        logger.error(f"Gemini Chat API error: {e}")
        if "429" in str(e): return "⚠ ERROR: Chat API quota exceeded. Please try again later."
        if "404" in str(e): return "⚠ ERROR: Chat model not available."
        return f"⚠ ERROR: {e}"


def generate_image_description(image_path):
    try:
        if not gemini_client:
            return "⚠ ERROR: Gemini client not initialized. Please check GEMINI_API_KEY."
        ext = (os.path.splitext(image_path)[1] or "").lower()
        mime_type = "image/png" if ext == ".png" else "image/jpeg"
        with open(image_path, "rb") as img_file:
            img_data = img_file.read()
        image_part = genai_types.Part.from_bytes(data=img_data, mime_type=mime_type)
        content = genai_types.Content(
            role="user",
            parts=[
                genai_types.Part.from_text(text="Provide a detailed description of this image for educational purposes."),
                image_part,
            ],
        )
        response = gemini_client.models.generate_content(
            model=GEMINI_MODEL,
            contents=content,
        )
        text = getattr(response, "text", None) if response else None
        if not text or not text.strip():
            return "⚠ ERROR: No description returned."
        return text.strip().replace("*", "")
    except Exception as e:
        logger.error(f"Image description error: {e}")
        return f"⚠ ERROR: Unable to describe image - {e}"


@retry(stop=stop_after_attempt(2), retry=retry_if_exception_type(RequestsConnectionError))
def generate_flashcards_with_retry(extracted_text):
    try:
        clipped_text = extracted_text[:8000]
        prompt = f"""Generate up to 20 flashcards from the following text. Each flashcard must follow this exact format:
Question: [clear question text]
Answer: [concise answer text]

Ensure there is a blank line between flashcards. Do not use asterisks, numbers, or other markdown.
Ensure each question and answer is concise, non-empty, and relevant to the text.

Text to analyze:
{clipped_text}"""
        response = generate_gemini_response(prompt)
        if "⚠ ERROR" in response:
            return {"status": "error", "message": response.replace("⚠ ERROR: ", "")}
        return response
    except Exception as e:
        if "429" in str(e):
            return {"status": "error", "message": "Quota exceeded. Please try later."}
        raise


def process_flashcards(response_text):
    """Parse raw AI flashcard text into a structured list."""
    try:
        if isinstance(response_text, dict):
            if response_text.get("status") == "error":
                return []
            response_text = str(response_text)
        flashcards_list = [f.strip() for f in response_text.split("\n\n") if f.strip()]
        structured_flashcards = []
        for i, flashcard in enumerate(flashcards_list, 1):
            if i > 20: break
            parts = [p.strip() for p in flashcard.split("\n") if p.strip()]
            if len(parts) >= 2:
                front = back = None
                for part in parts:
                    if part.startswith("Question:") or part.startswith("Q:"):
                        front = part.split(":", 1)[1].strip()
                    elif part.startswith("Answer:") or part.startswith("A:"):
                        back = part.split(":", 1)[1].strip()
                if front and back and len(front) > 3 and len(back) > 3:
                    structured_flashcards.append({"id": i, "front": front, "back": back})
        return structured_flashcards
    except Exception as e:
        logger.error(f"Flashcard processing error: {e}")
        return []


@retry(stop=stop_after_attempt(2), retry=retry_if_exception_type(RequestsConnectionError))
def generate_short_notes_with_retry(extracted_text):
    try:
        clipped_text = extracted_text[:6000]
        prompt = f"""Generate concise short notes from the following text.
Use bullet points for key ideas and keep each point brief (1-2 sentences).
Ensure clarity and relevance for study purposes.
Do not use markdown symbols like asterisks.

Text:
{clipped_text}"""
        response = generate_gemini_response(prompt)
        if "⚠ ERROR" in response:
            return {"status": "error", "message": response.replace("⚠ ERROR: ", "")}
        return response
    except Exception as e:
        if "429" in str(e):
            return {"status": "error", "message": "Quota exceeded. Please try later."}
        raise


@retry(stop=stop_after_attempt(3), wait=wait_fixed(2))
def generate_mcqs_with_retry(extracted_text, num_questions: int = 10):
    """Generate MCQs with retry logic and error handling."""
    response = ""
    try:
        prompt = (
            f"Create {num_questions} multiple-choice questions from the following content. "
            'Return ONLY a valid JSON array with this exact format:\n'
            '[{"question": "Question text", "options": ["Option A", "Option B", "Option C", "Option D"], "answer": "Correct option text"}]\n'
            "Ensure each question is clear, options are distinct, and the answer matches exactly one option.\n\n"
            "Content:\n" + extracted_text[:8000]
        )
        response = generate_gemini_response(prompt)
        if "ERROR" in response:
            raise Exception("API error occurred")

        raw = response.strip()
        for prefix in ("```json", "```JSON", "```"):
            if raw.startswith(prefix):
                raw = raw[len(prefix):]
                break
        if raw.endswith("```"):
            raw = raw[:-3]
        raw = raw.strip()

        if not (raw.startswith("[") and raw.endswith("]")):
            start, end = raw.find("["), raw.rfind("]")
            if start != -1 and end != -1 and end > start:
                raw = raw[start:end + 1]

        data = json.loads(raw)
        cleaned = []
        for item in data[:num_questions]:
            if not isinstance(item, dict): continue
            q    = (item.get("question") or "").strip()
            opts = item.get("options") or []
            ans  = (item.get("answer") or "").strip()
            if q and isinstance(opts, list) and len(opts) >= 4 and ans:
                options_list = [str(opt).strip() for opt in opts[:4]]
                if ans in options_list:
                    cleaned.append({
                        "question": q.replace("**", "").replace("*", ""),
                        "options":  options_list,
                        "answer":   ans.replace("**", "").replace("*", ""),
                    })
        return cleaned if cleaned else [{"error": "No valid MCQs could be generated from the AI response."}]

    except json.JSONDecodeError as e:
        logger.error(f"JSON parsing error in MCQ generation: {e}\nRaw: {response}")
        return [{"error": "Failed to parse MCQs from AI response."}]
    except Exception as e:
        logger.error(f"MCQ generation error: {e}")
        if "429" in str(e):
            return [{"error": "MCQ generation quota exceeded. Please wait and try again."}]
        return [{"error": f"MCQ generation failed: {str(e)}"}]


def format_mcq_items(raw_mcqs):
    if isinstance(raw_mcqs, list) and len(raw_mcqs) == 1 and isinstance(raw_mcqs[0], dict) and "error" in raw_mcqs[0]:
        return []
    formatted = []
    for i, mcq in enumerate(raw_mcqs):
        if "error" in mcq: continue
        question = mcq.get('question', '').strip()
        options  = mcq.get('options', [])
        answer   = mcq.get('answer', '').strip()
        if question and len(options) >= 4 and answer:
            correct_idx = options.index(answer) if answer in options else 0
            formatted.append({
                "id":             i + 1,
                "question":       question,
                "options": [
                    {"letter": "A", "text": str(options[0]), "is_correct": correct_idx == 0},
                    {"letter": "B", "text": str(options[1]), "is_correct": correct_idx == 1},
                    {"letter": "C", "text": str(options[2]), "is_correct": correct_idx == 2},
                    {"letter": "D", "text": str(options[3]), "is_correct": correct_idx == 3},
                ],
                "answer":         answer,
                "correct_answer": ["A", "B", "C", "D"][correct_idx],
                "explanation":    f"The correct answer is {answer}.",
                "user_answer":    "",
                "is_answered":    False,
                "is_revealed":    False,
            })
    return formatted


@app.route('/upload', methods=['POST'])
def upload_file_or_url():
    """Handle file uploads and URL processing with optimised memory management."""
    logger.info("Received upload request")
    files_data = []
    temp_files = []
    quick_mode = request.form.get('quick_mode') in ['1', 'true', 'True']

    try:
        if 'files' not in request.files and 'youtube_url' not in request.form:
            return jsonify({"error": "No files or YouTube URL provided."}), 400

        if 'files' in request.files:
            files = request.files.getlist('files')
            if not files or all(f.filename == '' for f in files):
                return jsonify({"error": "No valid files uploaded."}), 400

            for file in files:
                if not file.filename:
                    continue
                if not file.filename.lower().endswith(('.png', '.jpg', '.jpeg', '.pdf')):
                    return jsonify({"error": f"Unsupported file type for {file.filename}."}), 400

                filepath = os.path.join(UPLOAD_FOLDER, secure_filename(file.filename))
                temp_files.append(filepath)
                try:
                    file.save(filepath)
                    file_data = {"filename": file.filename}

                    if file.filename.lower().endswith(('.png', '.jpg', '.jpeg')):
                        file_data["type"] = "image"
                        file_data["image_description"] = generate_image_description(filepath)
                        ext = file.filename.lower()
                        mime_type = "image/jpeg" if ext.endswith(('.jpg', '.jpeg')) else "image/png"
                        with open(filepath, "rb") as img_file:
                            img_bytes = img_file.read()
                        file_data["base64_image"] = f"data:{mime_type};base64,{base64.b64encode(img_bytes).decode('utf-8')}"
                        del img_bytes
                        files_data.append(file_data)

                    elif file.filename.lower().endswith('.pdf'):
                        file_data["type"] = "pdf"
                        max_pages = 50 if quick_mode else 100
                        extracted_text = extract_text_from_pdf(filepath, max_pages=max_pages)

                        if not extracted_text:
                            logger.info("No text found with PyMuPDF, trying OCR…")
                            ocr_pages = 10 if quick_mode else 20
                            extracted_text = extract_text_from_scanned_pdf(filepath, max_pages=ocr_pages)

                        file_data["extracted_text"] = extracted_text or "No text could be extracted from this PDF."

                        if rag_processor and extracted_text and "No text could be extracted" not in extracted_text:
                            try:
                                book_id = secure_filename(file.filename).replace('.pdf', '').replace(' ', '_')[:50]
                                logger.info(f"Processing {file.filename} into RAG vector store…")
                                chunk_count = rag_processor.process_document(
                                    file_path=filepath,
                                    book_id=book_id,
                                    metadata={"filename": file.filename, "type": "pdf"},
                                    max_pages=max_pages,
                                )
                                file_data["book_id"] = book_id
                                file_data["rag_processed"] = True
                                file_data["rag_chunks"] = chunk_count
                                logger.info(f"✓ Stored {chunk_count} RAG chunks for {file.filename}")
                            except Exception as rag_error:
                                logger.warning(f"RAG processing failed: {rag_error}")
                                file_data["rag_processed"] = False
                        else:
                            file_data["rag_processed"] = False

                        files_data.append(file_data)

                except Exception as e:
                    logger.error(f"Error processing file {file.filename}: {e}")
                    continue

        if 'youtube_url' in request.form:
            youtube_url = request.form['youtube_url']
            if not youtube_url:
                return jsonify({"error": "Empty YouTube URL provided."}), 400
            try:
                extracted_text, transcript_data = extract_text_from_youtube(youtube_url)
                if not extracted_text:
                    video_id = extract_video_id(youtube_url)
                    msg = "Failed to extract content from YouTube video. "
                    if video_id:
                        msg += f"The video (ID: {video_id}) may not have captions or may be private/restricted."
                    else:
                        msg += "Invalid YouTube URL format."
                    return jsonify({"error": msg}), 400

                files_data.append({
                    "type": "youtube",
                    "filename": youtube_url,
                    "extracted_text": extracted_text,
                    "transcript": transcript_data,
                    "youtube_id": extract_video_id(youtube_url),
                })
            except Exception as e:
                logger.error(f"Error processing YouTube URL: {e}")
                return jsonify({"error": f"Failed to process YouTube URL: {str(e)}"}), 500

        if not files_data:
            return jsonify({"error": "Failed to process any files or URLs."}), 400

        response_data = []
        for file_data in files_data:
            processed_data = {
                "type":       file_data["type"],
                "filename":   file_data["filename"],
                "summary":    "",
                "flashcards": [],
                "short_notes": "",
                "mcqs":       [],
                "raw_text":   "",
                "is_image":   file_data["type"] == "image",
            }

            if file_data.get("book_id"):
                processed_data["book_id"] = file_data["book_id"]
                processed_data["rag_processed"] = file_data.get("rag_processed", False)
                if file_data.get("rag_chunks"):
                    processed_data["rag_chunks"] = file_data["rag_chunks"]

            if file_data["type"] == "youtube":
                processed_data["transcript"]  = file_data.get("transcript")
                processed_data["youtube_id"]  = file_data.get("youtube_id")

            if file_data["type"] == "image":
                processed_data["image_description"] = file_data["image_description"]
                processed_data["base64_image"]       = file_data["base64_image"]

            elif file_data["type"] in ["pdf", "youtube"]:
                extracted_text = file_data["extracted_text"]
                if extracted_text and "No text could be extracted" not in extracted_text:
                    clipped = extracted_text[:4000] if quick_mode else extracted_text[:8000]
                    processed_data["raw_text"] = extracted_text

                    if quick_mode:
                        # Quick mode: only summary, no heavy tasks
                        processed_data["summary"] = generate_gemini_response(
                            f"Summarize this text concisely:\n\n{clipped}"
                        )
                    else:
                        # Run all 4 Gemini tasks in parallel — cuts 16-20s down to ~5s
                        def _summary():    return generate_gemini_response(f"Summarize this text concisely:\n\n{clipped}")
                        def _notes():      return generate_short_notes_with_retry(extracted_text)
                        def _flashcards(): return generate_flashcards_with_retry(extracted_text)
                        def _mcqs():       return generate_mcqs_with_retry(extracted_text, 10)

                        tasks = {
                            "summary":    _summary,
                            "notes":      _notes,
                            "flashcards": _flashcards,
                            "mcqs":       _mcqs,
                        }
                        results = {}
                        with ThreadPoolExecutor(max_workers=4) as pool:
                            futures = {pool.submit(fn): key for key, fn in tasks.items()}
                            for future in as_completed(futures):
                                key = futures[future]
                                try:
                                    results[key] = future.result()
                                except Exception as exc:
                                    logger.warning(f"{key} generation failed: {exc}")
                                    results[key] = None

                        processed_data["summary"] = results.get("summary") or ""

                        notes_resp = results.get("notes")
                        processed_data["short_notes"] = (
                            f"Note generation failed: {notes_resp.get('message')}"
                            if isinstance(notes_resp, dict) and notes_resp.get("status") == "error"
                            else (notes_resp or "")
                        )

                        fc_resp = results.get("flashcards")
                        processed_data["flashcards"] = (
                            [] if isinstance(fc_resp, dict) and fc_resp.get("status") == "error"
                            else process_flashcards(fc_resp) if fc_resp else []
                        )

                        processed_data["mcqs"] = format_mcq_items(results.get("mcqs") or [])

            response_data.append(processed_data)

        return jsonify(response_data)

    except Exception as e:
        logger.error(f"Unexpected error in upload_file_or_url: {e}")
        return jsonify({"error": "An unexpected error occurred while processing your request."}), 500

    finally:
        for filepath in temp_files:
            safe_remove(filepath)

@app.route('/generate/notes', methods=['POST'])
def generate_notes():
    data = request.get_json(force=True)
    text = data.get('text') or ''
    if not text:
        return jsonify({"error": "text is required"}), 400
    try:
        response = generate_short_notes_with_retry(text[:6000])
        if isinstance(response, dict) and response.get("status") == "error":
            return jsonify({"error": response.get("message")}), 500
        return jsonify({"short_notes": response})
    except Exception as e:
        logger.error(f"notes gen error: {e}")
        return jsonify({"error": str(e)}), 500


@app.route('/generate/flashcards', methods=['POST'])
def generate_flashcards_endpoint():
    data = request.get_json(force=True)
    text = data.get('text') or ''
    if not text:
        return jsonify({"error": "text is required"}), 400
    try:
        resp = generate_flashcards_with_retry(text[:8000])
        return jsonify({"flashcards": process_flashcards(resp)})
    except Exception as e:
        logger.error(f"flashcards gen error: {e}")
        return jsonify({"error": str(e)}), 500


@app.route('/generate/mcqs', methods=['POST'])
def generate_mcqs_endpoint():
    data = request.get_json() or {}
    text = data.get('text', '').strip()
    n    = min(max(1, int(data.get('num_questions', 10))), 20)

    if not text:
        return jsonify({"error": "Text is required"}), 400

    try:
        raw_mcqs = generate_mcqs_with_retry(text, n)

        if isinstance(raw_mcqs, list) and len(raw_mcqs) == 1 and isinstance(raw_mcqs[0], dict) and "error" in raw_mcqs[0]:
            return jsonify({"mcqs": [], "error": raw_mcqs[0]["error"]})

        formatted = []
        for i, mcq in enumerate(raw_mcqs):
            if "error" in mcq: continue
            question = mcq.get('question', '').strip()
            options  = mcq.get('options', [])
            answer   = mcq.get('answer', '').strip()
            if question and len(options) >= 4 and answer:
                correct_idx = options.index(answer) if answer in options else 0
                formatted.append({
                    "id":             i + 1,
                    "question":       question,
                    "options": [
                        {"letter": "A", "text": str(options[0]), "is_correct": correct_idx == 0},
                        {"letter": "B", "text": str(options[1]), "is_correct": correct_idx == 1},
                        {"letter": "C", "text": str(options[2]), "is_correct": correct_idx == 2},
                        {"letter": "D", "text": str(options[3]), "is_correct": correct_idx == 3},
                    ],
                    "answer":         answer,
                    "correct_answer": ["A", "B", "C", "D"][correct_idx],
                    "explanation":    f"The correct answer is {answer}.",
                    "user_answer":    "",
                    "is_answered":    False,
                    "is_revealed":    False,
                })

        return jsonify({"mcqs": formatted, "total_generated": len(formatted)})

    except Exception as e:
        logger.error(f"MCQ generation error: {e}")
        return jsonify({"error": "Failed to generate MCQs. Please try again."}), 500


@app.route('/mcq/submit', methods=['POST'])
def submit_mcq_answers():
    data = request.get_json()
    mcqs = data.get('mcqs', [])
    if not mcqs:
        return jsonify({"error": "No answers provided"}), 400

    correct = sum(1 for q in mcqs if q.get('user_answer') == q.get('correct_answer'))
    total   = len(mcqs)
    score   = (correct / total) * 100 if total else 0
    feedback = (
        "Excellent!"     if score >= 90 else
        "Great!"         if score >= 70 else
        "Good effort!"   if score >= 50 else
        "Keep practicing!"
    )
    return jsonify({"success": True, "score": round(score, 1), "correct": correct, "total": total, "feedback": feedback})


@app.route('/mcq/progress', methods=['GET'])
def get_mcq_progress():
    return jsonify({"total_quizzes_taken": 0, "average_score": 0, "best_score": 0, "quizzes_completed": 0, "improvement_trend": "beginner"})

@app.route('/download', methods=['POST'])
def download_content():
    data         = request.get_json()
    content_type = data.get('type')
    format_type  = data.get('format')
    content      = data.get('content')

    if content_type == 'transcript':
        transcript_data = content.get('transcript', [])
        if format_type == 'txt':
            buffer = BytesIO()
            buffer.write(b"YouTube Video Transcript\n\n")
            for segment in transcript_data:
                if not isinstance(segment, dict): continue
                start_time = segment.get('start', 0)
                time_str = f"[{int(start_time // 60)}:{int(start_time % 60):02d}]"
                buffer.write(f"{time_str} {segment.get('text', '').strip()}\n".encode('utf-8'))
            buffer.seek(0)
            return send_file(buffer, as_attachment=True, download_name='transcript.txt', mimetype='text/plain')

    if not content_type or not format_type or not content:
        return jsonify({"error": "Missing required fields"}), 400

    filename = f"{content_type}.{format_type}"
    buffer   = BytesIO()

    if format_type == 'pdf':
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

        doc    = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        custom_style = ParagraphStyle(
            'CustomStyle', parent=styles['Normal'],
            fontSize=12, leading=14, spaceAfter=12, alignment=4,
        )
        story = []
        if content_type in ['summary', 'short_notes', 'image_description']:
            story.append(Paragraph(escape(content_type.replace('_', ' ').title()), styles['Title']))
            for paragraph in str(content).split('\n\n'):
                story.append(Paragraph(escape(paragraph), custom_style))
                story.append(Spacer(1, 12))
        elif content_type == 'flashcards':
            story.append(Paragraph("Flashcards", styles['Title']))
            if isinstance(content, list):
                for i, card in enumerate(content, 1):
                    if not isinstance(card, dict): continue
                    story.append(Paragraph(f"Card {i}", styles['Heading2']))
                    story.append(Paragraph(f"Question: {escape(str(card.get('front', '')))}", custom_style))
                    story.append(Paragraph(f"Answer: {escape(str(card.get('back', '')))}", custom_style))
                    story.append(Spacer(1, 12))
        elif content_type == 'mcqs':
            story.append(Paragraph("Multiple Choice Questions", styles['Title']))
            if isinstance(content, list):
                for i, mcq in enumerate(content, 1):
                    if not isinstance(mcq, dict): continue
                    story.append(Paragraph(f"Question {i}: {escape(str(mcq.get('question', '')))}", styles['Heading2']))
                    for option in mcq.get('options', []):
                        if not isinstance(option, dict): continue
                        story.append(Paragraph(f"{escape(str(option.get('letter', '')))}) {escape(str(option.get('text', '')))}", custom_style))
                    story.append(Paragraph(f"Correct Answer: {escape(str(mcq.get('correct_answer', '')))}", custom_style))
                    story.append(Spacer(1, 12))
        if not story:
            story.append(Paragraph("No content available for export.", custom_style))
        doc.build(story)
        buffer.seek(0)
        return send_file(buffer, as_attachment=True, download_name=filename, mimetype='application/pdf')

    elif format_type == 'txt':
        if content_type in ['summary', 'short_notes', 'image_description']:
            buffer.write(content.encode('utf-8'))
        elif content_type == 'flashcards':
            buffer.write(b"Flashcards\n\n")
            for i, card in enumerate(content, 1):
                buffer.write(f"Card {i}\n  Question: {card['front']}\n  Answer: {card['back']}\n\n".encode('utf-8'))
        elif content_type == 'mcqs':
            buffer.write(b"Multiple Choice Questions\n\n")
            for i, mcq in enumerate(content, 1):
                buffer.write(f"Question {i}: {mcq['question']}\n".encode('utf-8'))
                for option in mcq['options']:
                    buffer.write(f"  {option['letter']}) {option['text']}\n".encode('utf-8'))
                buffer.write(f"Correct Answer: {mcq['correct_answer']}\n".encode('utf-8'))
                buffer.write(f"Explanation: {mcq.get('explanation', 'No explanation available')}\n\n".encode('utf-8'))
        buffer.seek(0)
        return send_file(buffer, as_attachment=True, download_name=filename, mimetype='text/plain')

    elif format_type == 'docx':
        from docx import Document as DocxDocument
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = DocxDocument()
        doc.add_heading(content_type.replace('_', ' ').title(), level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
        if content_type in ['summary', 'short_notes', 'image_description']:
            for paragraph in content.split('\n\n'):
                p = doc.add_paragraph(paragraph)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif content_type == 'flashcards':
            for i, card in enumerate(content, 1):
                doc.add_paragraph(f"Card {i}", style='Heading 2')
                doc.add_paragraph(f"Question: {card['front']}")
                doc.add_paragraph(f"Answer: {card['back']}", style='Normal').runs[0].italic = True
        elif content_type == 'mcqs':
            for i, mcq in enumerate(content, 1):
                doc.add_paragraph(f"Question {i}: {mcq['question']}", style='Heading 2')
                for option in mcq['options']:
                    doc.add_paragraph(f"{option['letter']}) {option['text']}")
                doc.add_paragraph(f"Correct Answer: {mcq['correct_answer']}", style='Normal').runs[0].bold = True
                doc.add_paragraph(f"Explanation: {mcq.get('explanation', 'No explanation available')}", style='Normal')
        doc.save(buffer)
        buffer.seek(0)
        return send_file(buffer, as_attachment=True, download_name=filename,
                         mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    return jsonify({"error": "Unsupported format"}), 400


@app.route('/generate/ppt', methods=['POST'])
def generate_ppt():
    return jsonify({
        "success": False,
        "error": "PowerPoint generation is currently under maintenance. Please use PDF or DOCX export for now."
    }), 501

@app.route('/chat', methods=['POST'])
def chat():
    try:
        data             = request.get_json()
        user_message     = data.get('message')
        conversation_history = data.get('history', [])
        content          = data.get('content', {})
        if not user_message:
            return jsonify({"error": "No message provided."}), 400

        def _truncate(txt, max_len):
            return txt[:max_len] if isinstance(txt, str) else ""

        sources = []
        rag_context = ""
        book_id = content.get('book_id')
        if rag_processor and book_id:
            try:
                rag_results = rag_processor.query_book(book_id, user_message, k=3)
                if rag_results.get('results'):
                    rag_context = "\n\nRelevant content from your document:\n"
                    for i, result in enumerate(rag_results['results'], 1):
                        snippet = _truncate(result.get('content', ''), 300)
                        page = result.get('page') or i
                        sources.append({"id": i, "page": page, "snippet": snippet})
                        rag_context += f"[{i}] {snippet}\n   (Page {page})\n"
                    logger.info(f"Retrieved {len(rag_results['results'])} RAG chunks")
            except Exception as rag_error:
                logger.warning(f"RAG query failed: {rag_error}")

        content_context = ""
        if content.get("summary"):
            content_context += f"Summary:\n{_truncate(content['summary'], 2000)}\n\n"
        if content.get("short_notes"):
            content_context += f"Short notes:\n{_truncate(content['short_notes'], 2000)}\n\n"
        if content.get("image_description"):
            content_context += f"Image description:\n{_truncate(content['image_description'], 1000)}\n\n"

        recent_history = conversation_history[-5:] if isinstance(conversation_history, list) else []
        history_text   = "\n".join(
            f"{m.get('sender', 'user')}: {_truncate(m.get('text', ''), 200)}"
            for m in recent_history if isinstance(m, dict)
        )

        prompt = (
            "You are a friendly and helpful study assistant chatbot. "
            "Answer questions about the user's uploaded study content and engage in general conversation. "
            "Use the provided content and conversation history for context-aware, natural, concise responses. "
            "Keep responses under 150 words and maintain a friendly tone.\n\n"
        )
        if rag_context:
            prompt += f"{rag_context}\n"
        prompt += (
            f"Uploaded content context:\n{content_context}\n\n"
            f"Conversation history:\n{history_text}\n\n"
            f"User message: {_truncate(user_message, 300)}\n\n"
            "Respond as the chatbot:"
        )

        response_text = generate_chat_response(prompt)
        if "⚠ ERROR" in response_text:
            fallback = (
                "I'm having trouble reaching the AI right now. "
                "Try again in a moment, or ask a simpler question."
            )
            return jsonify({"response": fallback, "sources": []})
        return jsonify({"response": response_text, "sources": sources})

    except Exception as e:
        logger.error(f"Chat endpoint error: {e}")
        return jsonify({"error": f"Chat error: {str(e)}"}), 500

@app.route('/debug/rag-status')
def debug_rag_status():
    """Lightweight RAG status – no heavy dynamic imports."""
    return jsonify({
        "rag_enabled":             ENABLE_RAG,
        "rag_processor_available": rag_processor is not None,
        "rag_backend":             "gemini-embeddings" if rag_processor else None,
        "import_error":            rag_import_error,
        "init_error":              rag_init_error,
        "vector_store_path":       os.path.join(os.path.dirname(os.path.abspath(__file__)), 'vector_store'),
        "vector_store_exists":     os.path.exists(os.path.join(os.path.dirname(os.path.abspath(__file__)), 'vector_store')),
    })


@app.route('/api/rag/books', methods=['GET'])
def list_rag_books():
    if not rag_processor:
        return jsonify({"error": "RAG processor not available"}), 503
    try:
        import chromadb
        client      = chromadb.PersistentClient(path=rag_processor.persist_directory)
        collections = client.list_collections()
        books = []
        for collection in collections:
            try:
                books.append(rag_processor.get_stats(collection.name))
            except Exception as e:
                books.append({"book_id": collection.name, "status": "error", "error": str(e)})
        return jsonify({"books": books, "total": len(books)})
    except Exception as e:
        logger.error(f"Error listing RAG books: {e}")
        return jsonify({"error": str(e)}), 500


@app.route('/api/rag/query', methods=['POST'])
def rag_query():
    if not rag_processor:
        return jsonify({"error": "RAG processor not available"}), 503
    data     = request.get_json()
    book_id  = data.get('book_id')
    question = data.get('question')
    k        = data.get('k', 3)
    if not book_id or not question:
        return jsonify({"error": "book_id and question are required"}), 400
    try:
        return jsonify(rag_processor.query_book(book_id, question, k=k))
    except Exception as e:
        logger.error(f"RAG query error: {e}")
        return jsonify({"error": str(e)}), 500

SESSIONS_DIR = os.path.join(os.path.dirname(__file__), 'sessions')
os.makedirs(SESSIONS_DIR, exist_ok=True)

_SESSION_STRIP_FIELDS = {"raw_text", "base64_image"}


def _slim_file_for_session(file_dict: dict) -> dict:
    """Return a copy of file_dict with heavy fields removed for session storage."""
    return {k: v for k, v in file_dict.items() if k not in _SESSION_STRIP_FIELDS}


@app.route('/api/sessions', methods=['GET'])
def list_sessions():
    try:
        sessions = []
        if os.path.exists(SESSIONS_DIR):
            for filename in os.listdir(SESSIONS_DIR):
                if not filename.endswith('.json'): continue
                session_path = os.path.join(SESSIONS_DIR, filename)
                try:
                    with open(session_path, 'r', encoding='utf-8') as f:
                        session_data = json.load(f)
                    stat = os.stat(session_path)
                    sessions.append({
                        'id':         filename.replace('.json', ''),
                        'name':       session_data.get('name', 'Untitled Session'),
                        'created_at': session_data.get('created_at', ''),
                        'updated_at': datetime.datetime.fromtimestamp(stat.st_mtime).isoformat(),
                        'file_count': len(session_data.get('uploadedFiles', [])),
                        'has_chat':   len(session_data.get('chatHistory', [])) > 0,
                    })
                except Exception as e:
                    logger.warning(f"Error reading session {filename}: {e}")
        sessions.sort(key=lambda x: x.get('updated_at', ''), reverse=True)
        return jsonify({'sessions': sessions})
    except Exception as e:
        logger.error(f"Error listing sessions: {e}")
        return jsonify({"error": str(e)}), 500


@app.route('/api/sessions', methods=['POST'])
def save_session():
    try:
        data         = request.get_json()
        session_name = data.get('name', f"Session {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}")
        uploaded_files = [
            _slim_file_for_session(f) if isinstance(f, dict) else f
            for f in data.get('uploadedFiles', [])
        ]
        chat_history = data.get('chatHistory', [])
        session_id   = f"session_{int(time.time())}"
        session_data = {
            'id':            session_id,
            'name':          session_name,
            'created_at':    datetime.datetime.now().isoformat(),
            'uploadedFiles': uploaded_files,
            'chatHistory':   chat_history,
        }
        session_path = os.path.join(SESSIONS_DIR, f"{session_id}.json")
        with open(session_path, 'w', encoding='utf-8') as f:
            json.dump(session_data, f, indent=2, ensure_ascii=False)
        logger.info(f"Session saved: {session_id}")
        return jsonify({'success': True, 'session_id': session_id, 'message': 'Session saved successfully'})
    except Exception as e:
        logger.error(f"Error saving session: {e}")
        return jsonify({"error": str(e)}), 500


@app.route('/api/sessions/<session_id>', methods=['GET'])
def get_session(session_id):
    try:
        session_path = os.path.join(SESSIONS_DIR, f"{session_id}.json")
        if not os.path.exists(session_path):
            return jsonify({"error": "Session not found"}), 404
        with open(session_path, 'r', encoding='utf-8') as f:
            return jsonify(json.load(f))
    except Exception as e:
        logger.error(f"Error loading session: {e}")
        return jsonify({"error": str(e)}), 500


@app.route('/api/sessions/<session_id>', methods=['DELETE'])
def delete_session(session_id):
    try:
        session_path = os.path.join(SESSIONS_DIR, f"{session_id}.json")
        if not os.path.exists(session_path):
            return jsonify({"error": "Session not found"}), 404
        os.remove(session_path)
        logger.info(f"Session deleted: {session_id}")
        return jsonify({'success': True, 'message': 'Session deleted successfully'})
    except Exception as e:
        logger.error(f"Error deleting session: {e}")
        return jsonify({"error": str(e)}), 500


if __name__ == '__main__':
    app.run(host="0.0.0.0", port=5000, debug=False)